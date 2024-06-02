import logging
from aiogram import Bot, Dispatcher, types
from aiogram.types import ParseMode
from aiogram.contrib.middlewares.logging import LoggingMiddleware
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.utils import executor
from docx import Document
from io import BytesIO

# Ваш токен, который вы получили от BotFather
API_TOKEN = '7133041801:AAGvWpIuFSBhrd1562Dpo8bmCDM_Q9PdBO8'

# Настройка логирования
logging.basicConfig(level=logging.INFO)

# Инициализация бота и диспетчера
bot = Bot(token=API_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(bot, storage=storage)
dp.middleware.setup(LoggingMiddleware())

# Определение состояний
class Form(StatesGroup):
    name = State()
    address = State()
    statement = State()

# Обработчик команды /start
@dp.message_handler(commands='start')
async def cmd_start(message: types.Message):
    await Form.name.set()
    await message.reply("Привет! Я помогу тебе создать заявление. Введи своё имя и фамилию.")

# Обработчик ввода имени и фамилии
@dp.message_handler(state=Form.name)
async def process_name(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['name'] = message.text
    await Form.next()
    await message.reply("Спасибо! Теперь введи свой адрес.")

# Обработчик ввода адреса
@dp.message_handler(state=Form.address)
async def process_address(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['address'] = message.text
    await Form.next()
    await message.reply("Спасибо! Теперь введи текст заявления.")

# Обработчик ввода текста заявления
@dp.message_handler(state=Form.statement)
async def process_statement(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['statement'] = message.text

        # Загрузка шаблона документа
        template_path = 'template.docx'
        doc = Document(template_path)

        # Замена заполнителей в документе
        replace_text(doc, '{{name}}', data['name'])
        replace_text(doc, '{{address}}', data['address'])
        replace_text(doc, '{{statement}}', data['statement'])

        # Сохранение документа в BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Отправка документа пользователю
        await message.reply_document(document=file_stream, filename='zayavlenie.docx')
        await message.reply("Заявление создано и отправлено!")

    # Завершение состояния
    await state.finish()

# Функция для замены текста в документе
def replace_text(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell, placeholder, replacement)

if __name__ == '__main__':
    executor.start_polling(dp, skip_updates=True)
